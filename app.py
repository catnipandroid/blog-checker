import os
import json
from io import BytesIO

import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from openai import OpenAI
import streamlit_authenticator as stauth

# =========================
# OpenAI 설정
# =========================

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client = None
if OPENAI_API_KEY:
    client = OpenAI(api_key=OPENAI_API_KEY)


# =========================
# 로그인 설정 (여기 해시값/이메일만 네 걸로 바꾸면 됨)
# =========================

credentials = {
    "usernames": {
        "jaehyun": {
            "name": "재현",
            "password": "$2b$12$여기에_bcrypt_해시_붙여넣기",  # make_hash.py 결과
            "email": "jaehyun@example.com",
            "roles": ["admin"],
        },
        # 팀원 추가 예시
        # "member1": {
        #     "name": "팀원1",
        #     "password": "$2b$12$팀원_해시값",
        #     "email": "member1@example.com",
        #     "roles": ["user"],
        # },
    }
}

authenticator = stauth.Authenticate(
    credentials,
    "blog_checker_cookie",   # cookie 이름
    "some_random_key_123",   # 시크릿 키 (임의 문자열)
    1,                       # 쿠키 만료일 (일)
)


# =========================
# 공통 유틸
# =========================

def highlight_paragraph(paragraph, color=WD_COLOR_INDEX.YELLOW):
    for run in paragraph.runs:
        run.font.highlight_color = color


def add_comment_below(doc: Document, paragraph, comment_text: str):
    new_para = doc.add_paragraph()
    run = new_para.add_run(f"[자동검수] {comment_text}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    for r in new_para.runs:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    paragraph._p.addnext(new_para._p)


def get_full_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


# =========================
# 1단계: 룰 기반 체크
# =========================

def check_utm_links(doc: Document, report: list):
    """http 들어갔는데 utm_ 없는 링크"""
    count = 0
    for para in doc.paragraphs:
        text = para.text
        if "http" in text and "utm_" not in text:
            highlight_paragraph(para)
            add_comment_below(doc, para, "UTM 파라미터가 누락되었습니다. (예: ?utm_source=...)")
            count += 1

    if count > 0:
        report.append(f"- [룰] UTM 누락 문단 {count}개")
    else:
        report.append("- [룰] UTM 관련 문제 없음")


def check_hashtags(doc: Document, report: list, config: dict):
    """권장 해시태그 포함 여부"""
    recommended = config["recommended_hashtags"]
    full_text = get_full_text(doc)
    missing = [t for t in recommended if t and t not in full_text]

    if missing:
        p = doc.add_paragraph()
        run = p.add_run("[자동검수] 아래 해시태그가 부족합니다: " + ", ".join(missing))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        for r in p.runs:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW

        report.append(f"- [룰] 해시태그 부족: {len(missing)}개 (권장 해시태그 일부 누락)")
    else:
        report.append("- [룰] 해시태그 모두 포함됨")


def check_shopby(doc: Document, report: list, config: dict):
    """샵바이 언급 문단"""
    shopby_keywords = config["shopby_keywords"]
    count = 0
    for para in doc.paragraphs:
        text = para.text
        if any(keyword.lower() in text.lower() for keyword in shopby_keywords):
            highlight_paragraph(para)
            add_comment_below(doc, para, "샵바이(Shopby) 관련 내용은 블로그에 포함될 수 없습니다.")
            count += 1

    if count > 0:
        report.append(f"- [룰] 샵바이 언급 문단 {count}개")
    else:
        report.append("- [룰] 샵바이 언급 없음")


def check_b2b_basic_feature(doc: Document, report: list, config: dict):
    """B2B + 기본 기능 뉘앙스"""
    b2b_keywords = config["b2b_keywords"]
    basic_keywords = config["basic_feature_keywords"]

    count = 0
    for para in doc.paragraphs:
        text = para.text
        if any(b in text for b in b2b_keywords) and any(k in text for k in basic_keywords):
            highlight_paragraph(para)
            add_comment_below(doc, para, "B2B 기능이 기본 제공된다는 오해를 줄 수 있는 표현입니다.")
            count += 1

    if count > 0:
        report.append(f"- [룰] B2B를 기본 기능처럼 표현한 문단 {count}개")
    else:
        report.append("- [룰] B2B 기본 기능 오해 표현 없음")


def check_haedream(doc: Document, report: list, config: dict):
    """해드림 언급 문단"""
    haedream_keywords = config["haedream_keywords"]
    count = 0
    for para in doc.paragraphs:
        text = para.text
        if any(k in text for k in haedream_keywords):
            highlight_paragraph(para)
            add_comment_below(doc, para, "해드림 표기 방식이 정책에 맞는지 확인이 필요합니다.")
            count += 1

    if count > 0:
        report.append(f"- [룰] 해드림 언급 문단 {count}개")
    else:
        report.append("- [룰] 해드림 언급 없음")


def check_media_count(doc: Document, report: list, min_images: int = 15):
    """이미지 개수 / 영상 URL 여부"""
    img_count = len(doc.inline_shapes)
    full_text = get_full_text(doc)

    # 아주 러프하게 영상 URL 체크
    has_video = any(k in full_text for k in ["youtube.com", "youtu.be", "vimeo.com", "video"])

    # 이미지 개수
    if img_count < min_images:
        p = doc.add_paragraph()
        add_comment_below(
            doc,
            p,
            f"이미지 개수가 부족합니다. (현재 {img_count}장 / 기준 {min_images}장 이상)"
        )
        report.append(f"- [룰] 이미지 개수 부족: {img_count}장 (기준 {min_images}장)")
    else:
        report.append(f"- [룰] 이미지 개수 충족: {img_count}장")

    # 영상
    if not has_video:
        report.append("- [룰] 동영상 삽입 없음 (영상 1개 이상 권장)")
    else:
        report.append("- [룰] 동영상 URL 포함됨 (youtube 등)")


def check_forbidden_terms(doc: Document, report: list, client_brands: list[str], competitors: list[str]):
    """고객사 브랜드 / 타사 금지어"""
    client_count = 0
    comp_count = 0

    for para in doc.paragraphs:
        text_lower = para.text.lower()
        if any(b.lower() in text_lower for b in client_brands):
            highlight_paragraph(para)
            add_comment_below(doc, para, "고객사 브랜드명 언급 금지 대상이 포함되어 있습니다.")
            client_count += 1

        if any(c.lower() in text_lower for c in competitors):
            highlight_paragraph(para)
            add_comment_below(doc, para, "타사(경쟁사) 언급이 포함되어 있습니다.")
            comp_count += 1

    report.append(f"- [룰] 고객사 브랜드 언급 문단: {client_count}개")
    report.append(f"- [룰] 타사/경쟁사 언급 문단: {comp_count}개")


def check_avoided_phrases(doc: Document, report: list, avoided_phrases: list[str]):
    """쇼핑몰호스팅사, 전자상거래 플랫폼, 반응형스킨 등 지양 표현"""
    count = 0
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        if any(p.lower() in text_lower for p in avoided_phrases):
            highlight_paragraph(para)
            add_comment_below(doc, para, "내부에서 지양하는 표현이 포함되어 있습니다. 문구 수정 필요.")
            count += 1

    report.append(f"- [룰] 지양 표현이 포함된 문단: {count}개")


def check_title_keyword(doc: Document, report: list, required_keyword: str | None):
    """제목에 필수 키워드 포함 여부"""
    if not required_keyword:
        report.append("- [룰] 제목 키워드 기준 미설정 (수동 체크)")
        return

    if not doc.paragraphs:
        report.append("- [룰] 문단이 없어 제목을 확인할 수 없음")
        return

    title_para = doc.paragraphs[0]
    if required_keyword not in title_para.text:
        highlight_paragraph(title_para)
        add_comment_below(
            doc,
            title_para,
            f"제목에 지정된 키워드('{required_keyword}')가 포함되어 있지 않습니다."
        )
        report.append("- [룰] 제목 키워드 미포함")
    else:
        report.append("- [룰] 제목에 지정 키워드 포함")


# =========================
# 2단계: LLM 기반 체크 (옵션)
# =========================

def analyze_paragraph_with_llm(text: str) -> dict | None:
    if not client:
        return None
    if not text.strip():
        return None

    prompt = f"""
너는 NHN커머스 고도몰 블로그 원고를 검수하는 어시스턴트다.

아래 문단을 보고 다음 항목들을 판단해라.
반드시 JSON 문자열만 출력하라.

규칙:
1) "b2b_as_basic":    B2B 기능이 기본 기능처럼 보이게 표현됐는지 여부.
2) "free_b2b_mix":    무료/0원 프로모션 + B2B 내용이 섞여 잘못된 뉘앙스를 주는지 여부.
3) "haedream_mislabel":  해드림을 공식 에이전시처럼 잘못 표기했는지 여부.
4) "typo_exists":     맞춤법/띄어쓰기 문제가 있는지 여부.
5) "typo_examples":   대표적 맞춤법 오류 단어 3개 이하.

출력 형식(JSON 예시):

{{
  "b2b_as_basic": false,
  "free_b2b_mix": true,
  "haedream_mislabel": false,
  "typo_exists": true,
  "typo_examples": ["예시1", "예시2"]
}}

검수할 문단:
\"\"\"{text}\"\"\"
"""

    try:
        resp = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
            timeout=20,
        )
    except Exception as e:
        print("[LLM 오류] 요청 중 예외 발생:", e)
        return None

    content = resp.output_text
    try:
        data = json.loads(content)
        return data
    except Exception:
        print("[LLM] JSON 파싱 실패. 응답:", content[:200], "...")
        return None


def check_with_llm(doc: Document, report: list, config: dict, use_llm: bool):
    if not use_llm or not client:
        if not client:
            report.append("- [LLM] OPENAI_API_KEY 미설정으로 LLM 검수는 수행되지 않았습니다.")
        else:
            report.append("- [LLM] LLM 검수 옵션이 꺼져 있습니다.")
        return

    suspicious_keywords = config["suspicious_keywords"]

    b2b_basic_count = 0
    free_b2b_mix_count = 0
    haedream_mislabel_count = 0
    typo_count = 0

    paragraphs = list(doc.paragraphs)
    total = len(paragraphs)

    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text or len(text) < 15:
            continue

        lower = text.lower()
        if not any(k.lower() in lower for k in suspicious_keywords):
            continue

        print(f"[LLM] {idx+1}/{total} 문단 검사 중...")

        result = analyze_paragraph_with_llm(text)
        if not result:
            continue

        if result.get("b2b_as_basic"):
            highlight_paragraph(para)
            add_comment_below(
                doc,
                para,
                "LLM: B2B 기능이 '기본 제공'처럼 보이는 표현입니다. "
                "커스터마이징이 필요하다는 점을 명시해야 합니다."
            )
            b2b_basic_count += 1

        if result.get("free_b2b_mix"):
            highlight_paragraph(para)
            add_comment_below(
                doc,
                para,
                "LLM: 무료/0원 프로모션과 B2B 튜닝 내용이 섞여, "
                "B2B도 무료로 시작 가능한 것처럼 보일 수 있습니다."
            )
            free_b2b_mix_count += 1

        if result.get("haedream_mislabel"):
            highlight_paragraph(para)
            add_comment_below(
                doc,
                para,
                "LLM: 해드림을 공식 에이전시/제작 대행사처럼 표현한 부분이 있습니다. "
                "‘맞춤 제작 상담을 통해 공식 에이전시를 연결’하는 역할로 표시해야 합니다."
            )
            haedream_mislabel_count += 1

        if result.get("typo_exists"):
            examples = result.get("typo_examples") or []
            example_text = ", ".join(examples) if examples else "대표적인 오류 예시를 확인해 주세요."
            add_comment_below(
                doc,
                para,
                f"LLM: 이 문단에 맞춤법/띄어쓰기/오탈자 문제가 있습니다. 예시: {example_text}"
            )
            typo_count += 1

    report.append(f"- [LLM] B2B 기본기능처럼 보이는 문단: {b2b_basic_count}개")
    report.append(f"- [LLM] 무료 프로모션과 B2B 튜닝이 혼용된 문단: {free_b2b_mix_count}개")
    report.append(f"- [LLM] 해드림 표기 오해 소지가 있는 문단: {haedream_mislabel_count}개")
    report.append(f"- [LLM] 맞춤법/오탈자 지적된 문단: {typo_count}개")


# =========================
# 한 파일 처리
# =========================

def process_docx(file, filename: str, config: dict, use_llm: bool):
    doc = Document(file)
    report: list[str] = []

    # 룰 기반
    check_media_count(doc, report, config["min_images"])
    check_utm_links(doc, report)
    check_hashtags(doc, report, config)
    check_shopby(doc, report, config)
    check_b2b_basic_feature(doc, report, config)
    check_haedream(doc, report, config)
    check_forbidden_terms(doc, report, config["client_brands"], config["competitor_keywords"])
    check_avoided_phrases(doc, report, config["avoided_phrases"])
    check_title_keyword(doc, report, config["title_required_keyword"])

    # LLM 기반
    check_with_llm(doc, report, config, use_llm)

    # 요약
    summary = doc.add_paragraph()
    summary_run = summary.add_run("[자동검수 요약]")
    summary_run.bold = True
    for line in report:
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, report


# =========================
# Streamlit UI + 로그인
# =========================

def main():
    st.set_page_config("블로그 원고 자동검수", layout="wide")

    # 1) 로그인 위젯 렌더
    authenticator.login(
        location="main",
        fields={
            "Form name": "로그인",
            "Username": "아이디",
            "Password": "비밀번호",
            "Login": "로그인",
        },
        key="Login",
    )

    # 2) 세션 상태에서 인증 결과 읽기
    auth_status = st.session_state.get("authentication_status", None)
    name = st.session_state.get("name", None)
    username = st.session_state.get("username", None)

    if auth_status is False:
        st.error("아이디 또는 비밀번호가 올바르지 않습니다.")
        return
    elif auth_status is None:
        st.info("아이디와 비밀번호를 입력해 주세요.")
        return

    # 로그인 성공
    authenticator.logout(button_name="로그아웃", location="sidebar", key="Logout")
    st.sidebar.markdown(f"**👤 {name}님 로그인 중**")

    st.title("📝 고도몰 블로그 원고 자동 검수 봇")
    st.markdown("워드(.docx) 원고를 업로드하면, 정책에 맞춰 자동으로 형광펜 + 코멘트를 달아줍니다.")

    # ---- 사이드바: 규칙 설정 ----
    with st.sidebar:
        st.header("⚙ 규칙 설정")

        min_images = st.number_input("최소 이미지 개수 기준", min_value=0, max_value=100, value=15, step=1)

        hashtags_input = st.text_area(
            "권장 해시태그 (쉼표로 구분)",
            "#자사몰제작,#자사몰만들기,#무료쇼핑몰만들기,#온라인쇼핑몰창업,#B2B몰제작",
            height=70,
        )
        recommended_hashtags = [h.strip() for h in hashtags_input.split(",") if h.strip()]

        b2b_input = st.text_area(
            "B2B 관련 키워드",
            "B2B,도매몰,도매 쇼핑몰,폐쇄몰,가맹점 발주,프랜차이즈",
            height=60,
        )
        b2b_keywords = [k.strip() for k in b2b_input.split(",") if k.strip()]

        basic_input = st.text_area(
            "‘기본 기능’ 뉘앙스 키워드",
            "기본 기능,기본기능,기본으로 제공,기본 탑재,별도 개발 없이,추가 개발 없이,바로 사용할 수 있는",
            height=70,
        )
        basic_feature_keywords = [k.strip() for k in basic_input.split(",") if k.strip()]

        shopby_input = st.text_area(
            "샵바이 관련 키워드",
            "샵바이,shopby,Shopby,SHOPBY,샵바이 엔터프라이즈",
            height=60,
        )
        shopby_keywords = [k.strip() for k in shopby_input.split(",") if k.strip()]

        haedream_input = st.text_area(
            "해드림 관련 키워드",
            "해드림,헤드림",
            height=50,
        )
        haedream_keywords = [k.strip() for k in haedream_input.split(",") if k.strip()]

        client_brands_input = st.text_area(
            "고객사 브랜드명 (언급 금지, 쉼표로)",
            "고객A,고객B",
            height=60,
        )
        client_brands = [c.strip() for c in client_brands_input.split(",") if c.strip()]

        competitors_input = st.text_area(
            "타사/경쟁사 키워드 (언급 금지, 쉼표로)",
            "카페24,아임웹,메이크샵,shopify",
            height=60,
        )
        competitor_keywords = [c.strip() for c in competitors_input.split(",") if c.strip()]

        avoided_input = st.text_area(
            "지양 표현 리스트 (쇼핑몰호스팅사, 전자상거래 플랫폼 등)",
            "쇼핑몰호스팅사,쇼핑몰 호스팅사,전자상거래 플랫폼,반응형 스킨,반응형스킨",
            height=70,
        )
        avoided_phrases = [p.strip() for p in avoided_input.split(",") if p.strip()]

        title_required_keyword = st.text_input(
            "제목에 반드시 들어가야 할 키워드 (없으면 비워두기)",
            "",
        )

        suspicious_input = st.text_area(
            "LLM 검수 대상 '의심 키워드'",
            "B2B,도매몰,폐쇄몰,프랜차이즈,가맹점,무료,0원,프로모션,무상,해드림,헤드림",
            height=70,
        )
        suspicious_keywords = [k.strip() for k in suspicious_input.split(",") if k.strip()]

        use_llm = st.checkbox("LLM 기반 뉘앙스/맞춤법 검수 사용", value=True)

        config = {
            "min_images": int(min_images),
            "recommended_hashtags": recommended_hashtags,
            "b2b_keywords": b2b_keywords,
            "basic_feature_keywords": basic_feature_keywords,
            "shopby_keywords": shopby_keywords,
            "haedream_keywords": haedream_keywords,
            "client_brands": client_brands,
            "competitor_keywords": competitor_keywords,
            "avoided_phrases": avoided_phrases,
            "title_required_keyword": title_required_keyword.strip() or None,
            "suspicious_keywords": suspicious_keywords,
        }

    # ---- 파일 업로드 ----
    uploaded_files = st.file_uploader(
        "검수할 워드 파일(.docx)을 업로드하세요 (여러 개 가능)",
        type=["docx"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        for file in uploaded_files:
            st.subheader(f"📄 {file.name}")
            buffer, report = process_docx(file, file.name, config, use_llm)

            st.text_area("검수 요약", "\n".join(report), height=180)

            st.download_button(
                "✅ 검수된 파일 다운로드",
                data=buffer,
                file_name=file.name.replace(".docx", "_checked.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


if __name__ == "__main__":
    main()
